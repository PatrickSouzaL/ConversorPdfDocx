"""Passo 3 — reconstrução do documento Word nativo.

Recria o fluxo do manual original inserindo, na ordem de leitura, parágrafos de
texto (nativo/editável) e as imagens recortadas dos prints — cada uma na sua
posição. O texto passa a ser 100% editável no Word (não é mais imagem).
"""

from __future__ import annotations

import io
from pathlib import Path

import cv2
import numpy as np
from docx import Document
from docx.shared import Emu

from .config import Block, BlockKind, PipelineConfig

_EMU_PER_INCH = 914_400


class DocxBuilder:
    """Acumula páginas e blocos em um único documento Word."""

    def __init__(self, cfg: PipelineConfig) -> None:
        self.cfg = cfg
        self.document = Document()
        section = self.document.sections[0]
        self._usable_width = int(section.page_width - section.left_margin - section.right_margin)
        self._first_page = True
        self.text_blocks = 0
        self.image_blocks = 0

    def add_page(self, bgr: np.ndarray, blocks: list[Block]) -> None:
        """Adiciona uma página (com quebra de página entre páginas)."""
        if not self._first_page:
            self.document.add_page_break()
        self._first_page = False

        for block in blocks:
            if block.kind is BlockKind.TEXT:
                self._add_text(block.text)
                self.text_blocks += 1
            else:
                self._add_image(bgr, block.bbox)
                self.image_blocks += 1

    def _add_text(self, text: str) -> None:
        """Insere o texto do bloco, um parágrafo por linha (preserva a estrutura)."""
        for line in text.split("\n"):
            line = line.strip()
            if line:
                self.document.add_paragraph(line)

    def _add_image(self, bgr: np.ndarray, bbox: tuple[int, int, int, int]) -> None:
        """Recorta a região da imagem ORIGINAL (colorida) e a insere no DOCX."""
        x, y, w, h = bbox
        crop = bgr[y : y + h, x : x + w]
        if crop.size == 0:
            return

        ok, buffer = cv2.imencode(".png", crop)
        if not ok:
            return

        native_width = int(w / self.cfg.dpi * _EMU_PER_INCH)
        width = min(native_width, self._usable_width)

        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(io.BytesIO(buffer.tobytes()), width=Emu(width))

    def save(self, path: Path) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(str(path))
